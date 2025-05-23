from tkinter import messagebox
from docx import Document
import re
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.oxml import OxmlElement
import os
import sys
from docx.enum.table import WD_TABLE_ALIGNMENT
from PIL import Image as PILImage
import win32com.client
import io
import tempfile
import time
from copy import deepcopy
from docx.text.paragraph import Paragraph

# Importo la funzione resource_path
try:
    from form_interface import resource_path
except ImportError:
    # Funzione di fallback se non possiamo importare da form_interface
    def resource_path(relative_path):
        """Ottiene il percorso assoluto per le risorse, funziona sia in modalità sviluppo che come file eseguibile"""
        try:
            # PyInstaller crea un cartella temporanea e memorizza il percorso in _MEIPASS
            base_path = sys._MEIPASS
        except Exception:
            base_path = os.path.abspath(".")
        
        return os.path.join(base_path, relative_path)

def generate_document(template_path=None, output_path=None, data=None, images=None):
    """
    Genera un documento Word basato su un modello, sostituendo i segnaposto con i dati forniti
    e inserendo le immagini indicate dove si trova il segnaposto {{foto}} o {{Foto}}.
    
    Args:
        template_path (str, optional): Percorso del file modello Word (.docx). Se None, usa il modello predefinito.
        output_path (str, optional): Percorso dove salvare il documento generato. Se None, salva nella directory dello script.
        data (dict, optional): Dizionario con i dati da inserire nel documento
        images (list, optional): Lista di dizionari con path, description e orientation delle immagini
    """
    # Debug: verifica i dati ricevuti all'inizio della funzione
    print("\nDEBUG - docx_generator.generate_document - Dati ricevuti:")
    if "Oggetto del Sopralluogo" in data:
        print(f"data['Oggetto del Sopralluogo'] = {repr(data['Oggetto del Sopralluogo'])}")
        print("Newline all'inizio della funzione generate_document:", [i for i, char in enumerate(data['Oggetto del Sopralluogo']) if char == '\n'])
    
    # Imposta il modello predefinito se non specificato
    if template_path is None:
        template_path = resource_path("Modello Inspection.docx")
    
    # Imposta il percorso di output predefinito se non specificato
    if output_path is None:
        script_dir = os.path.dirname(os.path.abspath(__file__))
        output_path = os.path.join(script_dir, "documento_generato.docx")
    
    # Converti i percorsi in percorsi assoluti
    template_path = os.path.abspath(template_path)
    output_path = os.path.abspath(output_path)
    
    # Inizializza data se non specificato
    if data is None:
        data = {}
    
    try:
        # Fase 1: genera il documento con python-docx per sostituire i segnaposti normali e aggiungere le immagini
        doc = Document(template_path)
        
        # Definizione pattern per i segnaposto
        placeholder_pattern = re.compile(r'\{\{([^}]+)\}\}')
        
        # Flag per tracciare se abbiamo già inserito le immagini
        images_inserted = False
        
        # Funzione per sostituire i segnaposti in un paragrafo
        def process_paragraph(paragraph):
            nonlocal images_inserted
            if any(placeholder in paragraph.text for placeholder in ["{{foto}}", "{{Foto}}"]) and images and len(images) > 0:
                parent = paragraph._p.getparent()
                index = list(parent).index(paragraph._p)
                parent.remove(paragraph._p)
                
                table = doc.add_table(rows=1, cols=1)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                set_table_border(table, False)
                
                tbl = table._tbl
                parent.insert(index, tbl)
                
                populate_images_table(doc, table, images)
                
                images_inserted = True
                return True
            else:
                replace_text_in_paragraph(paragraph, data)
                return False
        
        # Sostituzione dei segnaposto nelle tabelle
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        process_paragraph(paragraph)
        
        # Sostituzione dei segnaposto nel corpo del documento
        paragraphs = list(doc.paragraphs)
        for paragraph in paragraphs:
            process_paragraph(paragraph)
        
        # Sostituzione dei segnaposto negli header e footer
        for section in doc.sections:
            for header in section.header.paragraphs:
                process_paragraph(header)
            
            for footer in section.footer.paragraphs:
                process_paragraph(footer)
        
        # Se non abbiamo ancora inserito le immagini e ci sono immagini da inserire
        if not images_inserted and images and len(images) > 0:
            doc.add_heading('Documentazione Fotografica', level=1)
            insert_images_table_at_end(doc, images)
        
        # Salva il documento temporaneo
        temp_path = output_path + "_temp.docx"
        doc.save(temp_path)
        
        # Fase 2: usa win32com per gestire i checkbox
        # Importante: Crea una nuova istanza di Word anziché usare quella eventualmente già aperta
        try:
            # Inizializza COM
            import pythoncom
            pythoncom.CoInitialize()
            
            # Crea una nuova istanza di Word (evita di riutilizzare istanze già aperte)
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False  # Nascondi l'applicazione
            
            print(f"Apertura documento: {temp_path}")
            doc = word.Documents.Open(temp_path)
            
            print("\nGestione dei checkbox nel documento:")
            
            # Per ogni campo nei dati che è un booleano (checkbox)
            for field_name, value in data.items():
                if isinstance(value, bool):
                    # Costruisci il segnaposto
                    placeholder = f"{{{{{field_name}}}}}"
                    print(f"\nRicerca segnaposto: {placeholder}")
                    
                    # Configura la ricerca
                    find = word.Selection.Find
                    find.ClearFormatting()
                    find.Text = placeholder
                    find.Forward = True
                    find.Wrap = 1  # wdFindContinue
                    find.Format = False
                    find.MatchCase = True
                    find.MatchWholeWord = False
                    
                    # Torna all'inizio del documento
                    word.Selection.HomeKey(6)  # 6 = wdStory
                    
                    # Cerca il segnaposto
                    found = find.Execute()
                    if found:
                        print(f"Trovato segnaposto {placeholder} alla posizione {word.Selection.Start}")
                        placeholder_pos = word.Selection.Start
                        
                        # Cerca la checkbox più vicina PRIMA del segnaposto
                        closest_checkbox = None
                        min_distance = float('inf')
                        
                        for i in range(1, doc.FormFields.Count + 1):
                            field = doc.FormFields.Item(i)
                            if field.Type == 71:  # 71 = wdFieldFormCheckBox
                                # Calcola la distanza solo se la checkbox è PRIMA del segnaposto
                                if field.Range.Start < placeholder_pos:
                                    distance = placeholder_pos - field.Range.Start
                                    print(f"\nCheckbox #{i}:")
                                    print(f"  Posizione: {field.Range.Start}")
                                    print(f"  Distanza dal segnaposto: {distance} caratteri")
                                    
                                    if distance < min_distance:
                                        min_distance = distance
                                        closest_checkbox = field
                        
                        if closest_checkbox and min_distance < 100:  # Limita la distanza a 100 caratteri
                            print(f"\nTrovata checkbox più vicina a {placeholder} (distanza: {min_distance} caratteri)")
                            print(f"Impostazione valore a: {value}")
                            closest_checkbox.CheckBox.Value = True if value else False
                        
                        # Rimuovi le parentesi graffe dal segnaposto
                        word.Selection.Text = field_name
                        print(f"Segnaposto sostituito con: '{field_name}'")
                    else:
                        print(f"Segnaposto {placeholder} non trovato nel documento")
            
            # Salva il documento e chiudi SOLO quello che abbiamo aperto
            print(f"\nSalvataggio documento: {output_path}")
            doc.SaveAs(output_path)
            doc.Close()
            
            # Chiudi solo l'istanza di Word che abbiamo creato
            word.Quit()
            
            # Rilascia le risorse COM
            pythoncom.CoUninitialize()
            
        except Exception as e:
            print(f"Errore durante la gestione di Word: {str(e)}")
            # In caso di errore, prova a rilasciare le risorse
            try:
                if 'doc' in locals() and doc:
                    doc.Close(SaveChanges=False)
                if 'word' in locals() and word:
                    word.Quit()
                pythoncom.CoUninitialize()
            except:
                pass
            raise
        
        # Rimuovi il file temporaneo
        if os.path.exists(temp_path):
            os.remove(temp_path)
            
    except Exception as e:
        print(f"Errore durante la modifica del documento: {str(e)}")
        # In caso di errore, usa il documento temporaneo come output finale
        if 'temp_path' in locals() and os.path.exists(temp_path):
            if os.path.exists(output_path):
                os.remove(output_path)
            os.rename(temp_path, output_path)
    
    return output_path

def replace_text_in_paragraph(paragraph, data):
    """Sostituisce i segnaposto nel testo del paragrafo"""
    # DEBUG - Verifica input
    print("\nDEBUG - replace_text_in_paragraph:")
    print(f"Paragrafo: '{paragraph.text}'")
    print(f"Dati disponibili: {list(data.keys())}")
    
    # Cerca tutti i segnaposto nel testo del paragrafo
    placeholder_pattern = re.compile(r'\{\{([^}]+)\}\}')
    matches = list(placeholder_pattern.finditer(paragraph.text))
    
    if not matches:
        return paragraph
    
    # Funzione per normalizzare il testo (rimuove spazi, caratteri speciali e converte in minuscolo)
    def normalize_text(text):
        # Rimuove spazi multipli e converte in minuscolo
        normalized = ' '.join(text.lower().split())
        # Rimuove caratteri speciali mantenendo lettere, numeri e spazi
        normalized = ''.join(c for c in normalized if c.isalnum() or c.isspace())
        return normalized
    
    # Prima gestisci i checkbox
    for match in matches:
        placeholder = match.group(1)
        placeholder_text = f"{{{{{placeholder}}}}}"
        
        if placeholder.startswith("checkbox_"):
            field_name = placeholder[9:]  # Rimuove "checkbox_"
            value = data.get(field_name, False)
            
            # Trova il run contenente il placeholder
            for run in paragraph.runs:
                if placeholder_text in run.text:
                    run.text = run.text.replace(placeholder_text, placeholder)
                    break
    
    # Poi gestisci gli altri segnaposto
    for match in matches:
        placeholder = match.group(1)
        placeholder_text = f"{{{{{placeholder}}}}}"
        
        # Salta i segnaposto che sono checkbox o che sono usati per i checkbox
        if placeholder.startswith("checkbox_") or placeholder in ["Visivo", "Rilievo/Verifica misure", "Test/Collaudo", "Altro", 
                                                                  "Conforme/Positivo", "Non conforme", "Osservazione", 
                                                                  "D.L. Generale", "D.L. Strutture", "D.L. Facciate", 
                                                                  "D.L. Imp. Elettrici/Speciali", "D.L. Imp. Meccanici"]:
            continue
        
        # Normalizza il placeholder
        normalized_placeholder = normalize_text(placeholder)
        
        # Debug essenziale per i segnaposto
        print(f"\nSegnaposto trovato: '{placeholder}'")
        print(f"Chiavi disponibili nel dizionario: {list(data.keys())}")
        
        # Cerca il valore nel dizionario data con varie strategie
        value = None
        
        # 1. Prova con il placeholder esatto
        if placeholder in data:
            value = data[placeholder]
            print(f"✓ Trovato con placeholder esatto")
        
        # 2. Prova con il placeholder senza spazi
        if value is None and placeholder.replace(" ", "") in data:
            value = data[placeholder.replace(" ", "")]
            print(f"✓ Trovato con placeholder senza spazi")
        
        # 3. Prova con il placeholder normalizzato
        if value is None and normalized_placeholder in data:
            value = data[normalized_placeholder]
            print(f"✓ Trovato con placeholder normalizzato")
        
        # 4. Prova cercando nel dizionario normalizzato
        if value is None:
            normalized_data = {normalize_text(k): v for k, v in data.items()}
            if normalized_placeholder in normalized_data:
                value = normalized_data[normalized_placeholder]
                print(f"✓ Trovato nel dizionario normalizzato")
        
        # Se ancora non trovato, usa stringa vuota
        if value is None:
            value = ""
            print(f"✗ Nessun valore trovato per il placeholder")
        
        print(f"Valore finale: '{value}'")
        
        # Verifica se il valore contiene tag HTML (formattazione)
        if placeholder == "Oggetto del Sopralluogo" and ("<b>" in str(value) or "<i>" in str(value) or "<u>" in str(value)):
            # Trova il run contenente il placeholder
            target_run = None
            for run in paragraph.runs:
                if placeholder_text in run.text:
                    target_run = run
                    break
            
            if target_run:
                # Rimuovi il placeholder
                start_pos = target_run.text.find(placeholder_text)
                end_pos = start_pos + len(placeholder_text)
                text_before = target_run.text[:start_pos]
                text_after = target_run.text[end_pos:]
                
                # Imposta il testo prima del placeholder
                target_run.text = text_before
                
                # Analizza il testo con formattazione
                current_run = target_run
                is_bold = False
                is_italic = False
                is_underline = False
                text_buffer = ""
                
                i = 0
                while i < len(value):
                    # Gestione tag di apertura
                    if value[i:i+3] == "<b>":
                        # Crea un run con il buffer corrente
                        if text_buffer:
                            if current_run.text:
                                current_run = paragraph.add_run(text_buffer)
                                current_run.font.name = "Arial"
                                current_run.font.size = Pt(10)
                                if is_italic:
                                    current_run.italic = True
                                if is_underline:
                                    current_run.underline = True
                            else:
                                current_run.text = text_buffer
                            text_buffer = ""
                        is_bold = True
                        i += 3
                    elif value[i:i+3] == "<i>":
                        # Crea un run con il buffer corrente
                        if text_buffer:
                            if current_run.text:
                                current_run = paragraph.add_run(text_buffer)
                                current_run.font.name = "Arial"
                                current_run.font.size = Pt(10)
                                if is_bold:
                                    current_run.bold = True
                                if is_underline:
                                    current_run.underline = True
                            else:
                                current_run.text = text_buffer
                            text_buffer = ""
                        is_italic = True
                        i += 3
                    elif value[i:i+3] == "<u>":
                        # Crea un run con il buffer corrente
                        if text_buffer:
                            if current_run.text:
                                current_run = paragraph.add_run(text_buffer)
                                current_run.font.name = "Arial"
                                current_run.font.size = Pt(10)
                                if is_bold:
                                    current_run.bold = True
                                if is_italic:
                                    current_run.italic = True
                            else:
                                current_run.text = text_buffer
                            text_buffer = ""
                        is_underline = True
                        i += 3
                    # Gestione tag di chiusura
                    elif value[i:i+4] == "</b>":
                        # Crea un run in grassetto con il buffer corrente
                        if text_buffer:
                            current_run = paragraph.add_run(text_buffer)
                            current_run.font.name = "Arial"
                            current_run.font.size = Pt(10)
                            current_run.bold = True
                            if is_italic:
                                current_run.italic = True
                            if is_underline:
                                current_run.underline = True
                            text_buffer = ""
                        is_bold = False
                        i += 4
                    elif value[i:i+4] == "</i>":
                        # Crea un run in corsivo con il buffer corrente
                        if text_buffer:
                            current_run = paragraph.add_run(text_buffer)
                            current_run.font.name = "Arial"
                            current_run.font.size = Pt(10)
                            current_run.italic = True
                            if is_bold:
                                current_run.bold = True
                            if is_underline:
                                current_run.underline = True
                            text_buffer = ""
                        is_italic = False
                        i += 4
                    elif value[i:i+4] == "</u>":
                        # Crea un run sottolineato con il buffer corrente
                        if text_buffer:
                            current_run = paragraph.add_run(text_buffer)
                            current_run.font.name = "Arial"
                            current_run.font.size = Pt(10)
                            current_run.underline = True
                            if is_bold:
                                current_run.bold = True
                            if is_italic:
                                current_run.italic = True
                            text_buffer = ""
                        is_underline = False
                        i += 4
                    else:
                        # Aggiungi il carattere al buffer
                        text_buffer += value[i]
                        i += 1
                
                # Aggiungi qualsiasi testo rimasto nel buffer
                if text_buffer:
                    current_run = paragraph.add_run(text_buffer)
                    current_run.font.name = "Arial"
                    current_run.font.size = Pt(10)
                    if is_bold:
                        current_run.bold = True
                    if is_italic:
                        current_run.italic = True
                    if is_underline:
                        current_run.underline = True
                
                # Aggiungi il testo dopo il placeholder
                if text_after:
                    final_run = paragraph.add_run(text_after)
                    final_run.font.name = "Arial"
                    final_run.font.size = Pt(10)
        else:
            # Sostituzione normale per i campi senza formattazione
            # Dobbiamo gestire due casi:
            # 1. Il segnaposto è interamente in un run
            # 2. Il segnaposto è diviso tra più run
            
            # Caso 1: il segnaposto è in un singolo run
            replaced = False
            for run in paragraph.runs:
                if placeholder_text in run.text:
                    # Assicurati che i newline vengano preservati nel documento Word
                    replacement_text = str(value)
                    
                    # Debug: verifica la presenza di newline nel testo di sostituzione
                    if placeholder == "Oggetto del Sopralluogo":
                        print("\nDEBUG - Testo da inserire nel documento per 'Oggetto del Sopralluogo':")
                        print(f"Testo da inserire (len: {len(replacement_text)}):")
                        print(repr(replacement_text))  # Usa repr per mostrare i caratteri speciali come \n
                        print("Caratteri newline trovati alle posizioni:")
                        newline_positions = [i for i, char in enumerate(replacement_text) if char == '\n']
                        print(newline_positions)
                        
                        # Gestione speciale per Oggetto del Sopralluogo per forzare nuovi paragrafi
                        if '\n' in replacement_text:
                            print("TROVATI NEWLINE IN OGGETTO DEL SOPRALLUOGO - CREAZIONE PARAGRAFI")
                            
                            # Dividi il testo in righe
                            lines = replacement_text.split('\n')
                            print(f"Numero di righe dopo split: {len(lines)}")
                            
                            # Sostituisci il segnaposto con la prima riga
                            run.text = run.text.replace(placeholder_text, lines[0])
                            
                            # Ottieni il paragrafo attuale e il suo contenitore
                            current_p = paragraph._p
                            parent = current_p.getparent()
                            
                            try:
                                # Per ciascuna riga aggiuntiva, crea un vero nuovo paragrafo
                                for i in range(1, len(lines)):
                                    if not lines[i].strip():  # Salta righe vuote
                                        continue
                                    
                                    print(f"Creazione nuovo paragrafo per riga {i}: '{lines[i]}'")
                                    
                                    # Crea un nuovo paragrafo copiando il paragrafo originale
                                    # Questo mantiene tutta la formattazione e stile
                                    new_p = deepcopy(paragraph._p)
                                    
                                    # Rimuovi tutti i run esistenti dal paragrafo copiato
                                    for old_run in new_p.findall('.//w:r', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                                        old_run.getparent().remove(old_run)
                                    
                                    # Crea un nuovo run con le stesse proprietà del run originale
                                    new_r = deepcopy(run._r)
                                    # Rimuovi gli elementi di testo esistenti
                                    for old_t in new_r.findall('.//w:t', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                                        old_t.getparent().remove(old_t)
                                    
                                    # Aggiungi nuovo testo con la stessa formattazione
                                    new_t = OxmlElement('w:t')
                                    new_t.set(qn('xml:space'), 'preserve')  # Preserva spazi
                                    new_t.text = lines[i]
                                    new_r.append(new_t)
                                    
                                    # Aggiungi il run al paragrafo
                                    new_p.append(new_r)
                                    
                                    # Aggiungi il nuovo paragrafo dopo il paragrafo corrente
                                    current_p.addnext(new_p)
                                    # Aggiorna il paragrafo corrente per il prossimo ciclo
                                    current_p = new_p
                            
                            except Exception as e:
                                print(f"Errore durante la creazione dei paragrafi: {str(e)}")
                                # In caso di errore, torna al metodo originale
                                fallback_text = "\n".join(lines[1:])
                                fallback_run = paragraph.add_run("\n" + fallback_text)
                                fallback_run.font.name = "Arial"
                                fallback_run.font.size = Pt(10)
                            
                            replaced = True
                            break
                        else:
                            # Se non ci sono newline, sostituisci normalmente
                            run.text = run.text.replace(placeholder_text, replacement_text)
                            replaced = True
                            break

            # Caso 2: il segnaposto è diviso tra più run
            if not replaced:
                # Ottieni tutti i run che contengono parte del segnaposto
                run_parts = []
                start_idx = -1
                end_idx = -1
                
                # Trova gli indici dei run che contengono parti del segnaposto
                for i, run in enumerate(paragraph.runs):
                    if "{{" in run.text and start_idx == -1:
                        start_idx = i
                    if "}}" in run.text and start_idx != -1:
                        end_idx = i
                        break
                
                # Se abbiamo trovato l'inizio e la fine del segnaposto
                if start_idx != -1 and end_idx != -1:
                    # Testo prima del segnaposto (nel primo run)
                    start_run = paragraph.runs[start_idx]
                    pre_text = start_run.text.split("{{")[0]
                    
                    # Testo dopo il segnaposto (nell'ultimo run)
                    end_run = paragraph.runs[end_idx]
                    post_text = end_run.text.split("}}")[1] if "}}" in end_run.text else ""
                    
                    # Preparazione del testo di sostituzione
                    replacement_text = str(value)
                    
                    # Se ci sono newline nel testo di sostituzione
                    if '\n' in replacement_text:
                        print("TROVATI NEWLINE IN SEGNAPOSTO DIVISO TRA PIÙ RUN")
                        
                        # Dividi il testo in righe
                        lines = replacement_text.split('\n')
                        
                        # Sostituisci il segnaposto con la prima riga
                        start_run.text = pre_text + lines[0]
                        
                        # Cancella i run intermedi e l'ultimo run
                        for i in range(start_idx + 1, end_idx + 1):
                            paragraph.runs[i].text = ""
                        
                        # Ottieni un riferimento al documento
                        doc = paragraph._parent
                        
                        # Aggiungi un nuovo paragrafo per ogni riga aggiuntiva
                        for i in range(1, len(lines)):
                            # Salta righe vuote
                            if not lines[i].strip():
                                continue
                                
                            print(f"Inserimento paragrafo per la riga {i}: '{lines[i]}'")
                            
                            # Crea un nuovo paragrafo
                            new_para = doc.add_paragraph(lines[i])
                            
                            # Ottieni l'elemento XML del nuovo paragrafo
                            new_p = new_para._p
                            
                            # Rimuovi il paragrafo dalla fine del documento
                            doc._body._body.remove(new_p)
                            
                            # Inseriscilo dopo il paragrafo corrente
                            paragraph._element.addnext(new_p)
                            
                            # Aggiorna il paragrafo corrente per il prossimo ciclo
                            paragraph = new_para
                    else:
                        # Senza newline, sostituisci normalmente
                        start_run.text = pre_text + replacement_text
                        
                        # Imposta l'ultimo run con il testo che segue il segnaposto
                        if start_idx != end_idx:
                            end_run.text = post_text
                        
                        # Cancella i run intermedi se ce ne sono
                        for i in range(start_idx + 1, end_idx + 1):
                            paragraph.runs[i].text = ""

    return paragraph

def replace_in_paragraph_with_formatting(paragraph, old_text, new_text):
    """
    Sostituisce il testo in un paragrafo mantenendo la formattazione.
    
    Args:
        paragraph: Paragrafo Word da modificare
        old_text (str): Testo da sostituire (segnaposto)
        new_text (str): Nuovo testo da inserire
    """
    # Per ogni run nel paragrafo
    for i, run in enumerate(paragraph.runs):
        if old_text in run.text:
            # Salva la formattazione originale
            original_font = run.font
            original_bold = run.bold
            original_italic = run.italic
            original_underline = run.underline
            
            # Sostituisci il testo mantenendo la formattazione
            run.text = run.text.replace(old_text, new_text)
            
            # Mantieni la formattazione
            run.bold = original_bold
            run.italic = original_italic
            run.underline = original_underline
            
            # Interrompi dopo la prima sostituzione
            break
        # Se il segnaposto è diviso tra più run, implementazione più complessa
        elif old_text.startswith(run.text) and i < len(paragraph.runs) - 1:
            # Gestione dei segnaposto che si estendono su più run
            # (implementazione complessa, qui omessa per semplicità)
            pass

def insert_images_table_at_paragraph(doc, paragraph_index, images):
    """
    Inserisce una tabella con immagini dopo un paragrafo specifico.
    
    Args:
        doc: Documento Word
        paragraph_index: Indice del paragrafo dopo cui inserire la tabella
        images: Lista di immagini da inserire
    """
    # Crea un nuovo paragrafo per la tabella
    table_paragraph = doc.add_paragraph()
    
    # Sposta il paragrafo nella posizione corretta
    # Questo è un workaround poiché python-docx non supporta direttamente
    # l'inserimento di paragrafi in posizioni specifiche
    p = table_paragraph._p
    doc._body._body.remove(p)
    doc._body._body.insert(paragraph_index + 1, p)
    
    # Inserisci la tabella nel documento
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table, False)
    
    # Rimuovi l'ultima tabella (quella appena creata in fondo)
    # e spostala nella posizione corretta
    tbl = table._tbl
    doc._body._body.remove(tbl)
    doc._body._body.insert(paragraph_index + 2, tbl)
    
    # Configura la tabella e inserisci le immagini
    populate_images_table(doc, table, images)

def insert_images_table_at_end(doc, images):
    """
    Inserisce una tabella con immagini alla fine del documento.
    
    Args:
        doc: Documento Word
        images: Lista di immagini da inserire
    """
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table, False)
    
    # Configura la tabella e inserisci le immagini
    populate_images_table(doc, table, images)

def insert_images_table(cell, images):
    """
    Inserisce una tabella con immagini in una cella.
    
    Args:
        cell: Cella della tabella in cui inserire le immagini
        images: Lista di immagini da inserire
    """
    # Crea una tabella nella cella
    table = cell.add_table(rows=1, cols=1)
    set_table_border(table, False)
    
    # Configura la tabella e inserisci le immagini
    populate_images_table(None, table, images)

def get_image_size(img_path, rotation=0):
    """
    Ottieni le dimensioni di un'immagine, considerando l'eventuale rotazione.
    
    Args:
        img_path: Percorso dell'immagine
        rotation: Angolo di rotazione (0, 90, 180, 270)
    
    Returns:
        Tuple: (larghezza, altezza)
    """
    try:
        with PILImage.open(img_path) as img:
            width, height = img.width, img.height
            # Se la rotazione è 90 o 270 gradi, scambia larghezza e altezza
            if rotation in [90, 270]:
                return height, width
            return width, height
    except Exception:
        return 0, 0

def populate_images_table(doc, table, images):
    """
    Popola una tabella con immagini e didascalie.
    """
    # Calcola le dimensioni per le immagini
    if doc:
        section = doc.sections[0]
        # Calcola la larghezza disponibile (pagina intera meno margini)
        page_width_cm = section.page_width.cm
        margin_left_cm = section.left_margin.cm
        margin_right_cm = section.right_margin.cm
        available_width_cm = page_width_cm - margin_left_cm - margin_right_cm
        
        # Converti in EMU (1 cm = 360000 EMU)
        available_width_emu = int(available_width_cm * 360000)
        
        # Altezza massima (metà pagina meno spazio per didascalia)
        page_height_cm = section.page_height.cm
        margin_top_cm = section.top_margin.cm
        margin_bottom_cm = section.bottom_margin.cm
        available_height_cm = ((page_height_cm - margin_top_cm - margin_bottom_cm - 2) / 2) * 0.8  # -2 per le didascalie
        available_height_emu = int(available_height_cm * 360000)
    else:
        # Se siamo in una cella, usiamo dimensioni di default
        available_width_emu = int(Cm(15).emu)  # pagina A4
        available_height_emu = int(Cm(10).emu)
    
    print(f"\nDimensioni disponibili:")
    print(f"Larghezza massima: {available_width_cm:.2f} cm ({available_width_emu} EMU)")
    print(f"Altezza massima: {available_height_cm:.2f} cm ({available_height_emu} EMU)")
    
    # Assicurati che la tabella abbia una sola colonna
    if len(table.columns) > 1:
        # Se la tabella ha più di una colonna, ridimensionala a una colonna
        print("Ridimensionamento tabella a una sola colonna")
        # Rimuoviamo tutte le righe esistenti
        for row in table.rows[:]:
            if hasattr(row, '_element'):
                row._element.getparent().remove(row._element)
        # Creiamo una nuova riga con una sola colonna
        table._tbl.append(OxmlElement('w:tr'))
        tr = table._tbl.tr_lst[-1]
        tr.append(OxmlElement('w:tc'))
    
    # Calcola il numero di righe necessarie
    num_rows = len(images)
    
    # Crea una tabella con il numero corretto di righe
    table.style = 'Table Grid'
    
    # Aggiungi le righe necessarie
    for _ in range(num_rows - 1):  # -1 perché la tabella già ha una riga
        table.add_row()
    
    # Processa tutte le immagini
    for i, img_info in enumerate(images):
        cell = table.cell(i, 0)
        
        try:
            # Usa l'immagine processata se disponibile, altrimenti carica dal file
            if "processed_image" in img_info:
                img = img_info["processed_image"]
                print(f"\nUsando immagine processata: {img.width}x{img.height}")
            else:
                img = PILImage.open(img_info["path"])
                print(f"\nCaricando immagine da file: {img_info['path']}")
                
                # Gestione orientamento EXIF
                try:
                    exif = img.getexif()
                    orientation = exif.get(274, 1)
                    print(f"Orientamento EXIF: {orientation}")
                    
                    if orientation == 3:
                        img = img.rotate(180, expand=True)
                    elif orientation == 6:
                        img = img.rotate(270, expand=True)
                    elif orientation == 8:
                        img = img.rotate(90, expand=True)
                except Exception as e:
                    print(f"Errore nella lettura EXIF: {str(e)}")
                
                # Applica la rotazione manuale se presente
                rotation = img_info.get("rotation", 0)
                if rotation != 0:
                    print(f"Applicazione rotazione manuale: {rotation}°")
                    img = img.rotate(rotation, expand=True)
            
            # Ottieni le dimensioni dopo tutte le rotazioni
            orig_width, orig_height = img.size
            print(f"Dimensioni dopo rotazioni: {orig_width}x{orig_height}")
            
            # Calcola il rapporto d'aspetto dell'immagine
            aspect_ratio = orig_width / orig_height
            print(f"Rapporto d'aspetto: {aspect_ratio:.2f}")
            
            # Usa sempre l'altezza come riferimento
            target_height_emu = available_height_emu
            target_width_emu = int(target_height_emu * aspect_ratio)
            
            # Se la larghezza calcolata supera quella disponibile, ricalcola partendo dalla larghezza
            if target_width_emu > available_width_emu:
                target_width_emu = available_width_emu
                target_height_emu = int(target_width_emu / aspect_ratio)
            
            print(f"Dimensioni target: {target_width_emu/360000:.2f}x{target_height_emu/360000:.2f} cm")
            
            # Calcola le dimensioni in pixel per il ridimensionamento
            scale_factor = 1  # Ridotto da 2 a 1 per ridurre il numero di pixel
            width_px = int((target_width_emu / 360000) * 96 * scale_factor)  # 96 DPI
            height_px = int((target_height_emu / 360000) * 96 * scale_factor)
            
            # Assicurati che le dimensioni siano almeno 1 pixel
            width_px = max(1, width_px)
            height_px = max(1, height_px)
            
            print(f"Dimensioni in pixel: {width_px}x{height_px}")
            
            # Crea un'immagine temporanea con le dimensioni corrette
            temp_path = f"{img_info['path']}_temp.jpg"
            
            # Ridimensiona e salva l'immagine con una qualità più bassa
            img = img.resize((width_px, height_px), PILImage.LANCZOS)
            img.save(temp_path, quality=75, dpi=(150, 150), optimize=True)
            
            # Paragrafo per l'immagine
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.space_after = Pt(24)  # Aggiunge 24pt di spazio dopo l'immagine
            run = p.add_run()
            
            # Aggiungi l'immagine al documento
            picture = run.add_picture(temp_path)
            
            # Imposta le dimensioni esatte dell'immagine nel documento
            picture.width = target_width_emu
            picture.height = target_height_emu
            
            # Rimuovi il file temporaneo
            os.remove(temp_path)
            
            # Aggiungi una riga vuota per la spaziatura
            cell.add_paragraph()
            
            # Debug per verificare se la didascalia è presente
            print(f"\nDati didascalia per l'immagine {i}:")
            print(f"Descrizione: '{img_info.get('description', '')}'")
            print(f"Rotazione: {img_info.get('rotation', 0)}")
            
            # Didascalia
            caption_text = img_info.get("description", "")
            print(f"Didascalia da inserire: '{caption_text}'")
            
            caption = cell.add_paragraph(caption_text)
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            format_caption(caption)
            
            # Aggiungi una riga vuota dopo la didascalia
            cell.add_paragraph()
            
        except Exception as e:
            # In caso di errore, aggiungi un messaggio di errore invece dell'immagine
            p = cell.paragraphs[0]
            p.text = f"Errore nel caricamento dell'immagine: {str(e)}"
            p.space_after = Pt(24)  # Aggiunge 24pt di spazio dopo il messaggio di errore
    
    # Imposta la larghezza della colonna della tabella
    table.columns[0].width = available_width_emu
    
    # Rimuovi i bordi della tabella
    set_table_border(table, False)

def format_caption(paragraph):
    """
    Formatta una didascalia con uno stile coerente.
    
    Args:
        paragraph: Paragrafo Word da formattare come didascalia
    """
    # Ottieni il testo attuale della didascalia
    current_text = paragraph.text.strip()
    
    # Determina il numero di figura in base alla posizione
    figure_number = 1
    parent = paragraph._p.getparent()
    if parent and parent.tag.endswith('tc'):  # Se siamo in una cella della tabella
        # Trova la riga della tabella
        tr = parent.getparent()
        if tr and tr.tag.endswith('tr'):
            # Trova tutte le righe precedenti nella tabella
            table = tr.getparent()
            if table and table.tag.endswith('tbl'):
                # Conta le righe precedenti che contengono immagini
                for row in table.findall('.//w:tr', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                    if row == tr:
                        break
                    # Verifica se la riga contiene un'immagine
                    if any(run.findall('.//w:drawing', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}) 
                          for run in row.findall('.//w:r', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})):
                        figure_number += 1
    
    # Crea il testo completo per la didascalia
    if not current_text:
        # Se non c'è una descrizione, usa solo il numero di figura
        paragraph.text = f"Figura {figure_number}"
    elif not current_text.startswith("Figura"):
        # Se c'è una descrizione ma non inizia con "Figura", aggiungi il numero
        paragraph.text = f"Figura {figure_number} - {current_text}"
    
    # Applica la formattazione indipendentemente dal contenuto
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(10)
    
    # Aggiungi spaziatura extra dopo la didascalia
    paragraph.space_after = Pt(24)  # Aggiunge 24pt di spazio dopo il paragrafo

def set_table_border(table, has_border):
    """
    Imposta o rimuove i bordi di una tabella.
    
    Args:
        table: Tabella Word da modificare
        has_border (bool): True per mostrare i bordi, False per nasconderli
    """
    tbl = table._tbl  # Accesso all'elemento XML sottostante
    for cell in tbl.iter_tcs():
        tcPr = cell.tcPr
        tcBorders = OxmlElement('w:tcBorders')
        
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            if has_border:
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')
            else:
                border.set(qn('w:val'), 'nil')
            
            tcBorders.append(border)
        
        tcPr.append(tcBorders)

def create_checkbox_control(paragraph, checked=False, insert_after=None):
    """
    Crea un controllo casella di controllo di Word.
    
    Args:
        paragraph: Paragrafo Word dove inserire il checkbox
        checked (bool): Se True, il checkbox sarà selezionato
        insert_after: Elemento XML dopo il quale inserire il checkbox
    """
    print(f"\nDEBUG - create_checkbox_control:")
    print(f"Creating checkbox with checked={checked}")
    
    # Crea un nuovo run per l'inizio del campo
    run_begin = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run_begin._r.append(fldChar1)
    
    # Crea un nuovo run per il testo del campo
    run_instr = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'FORMCHECKBOX'
    run_instr._r.append(instrText)
    
    # Crea un nuovo run per la fine del campo
    run_end = paragraph.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    if checked:
        fldChar2.set(qn('w:checked'), '1')
    run_end._r.append(fldChar2)
    
    # Se è specificato insert_after, sposta i run nella posizione corretta
    if insert_after is not None:
        insert_after.addnext(run_end._r)
        run_end._r.addprevious(run_instr._r)
        run_instr._r.addprevious(run_begin._r)
    
    print("Checkbox control created successfully")
    return run_end._r

def replace_placeholder(paragraph, placeholder, value):
    """Sostituisce un segnaposto con il valore corrispondente"""
    print(f"\nDEBUG - replace_placeholder:")
    print(f"Placeholder: {placeholder}")
    print(f"Value: {value} (type: {type(value)})")
    print(f"Paragraph text: {paragraph.text}")
    
    # Cerca il segnaposto nel testo del paragrafo
    placeholder_text = f"{{{{{placeholder}}}}}"
    print(f"Looking for: {placeholder_text}")
    
    # Verifica se il segnaposto è presente nel testo
    if placeholder_text in paragraph.text:
        print("Found placeholder in paragraph text")
        # Gestione speciale per i checkbox
        if placeholder.startswith("checkbox_"):
            print("Found checkbox placeholder")
            # Se il valore è un booleano, crea un vero checkbox
            if isinstance(value, bool):
                print(f"Creating checkbox with value: {value}")
                # Trova il run che contiene il segnaposto
                for run in paragraph.runs:
                    if placeholder_text in run.text:
                        print(f"Found placeholder in run: {run.text}")
                        # Sostituisci il segnaposto con "trovato"
                        run.text = run.text.replace(placeholder_text, "trovato")
                        print("Placeholder replaced with 'trovato'")
                        return
                # Se non troviamo il segnaposto nei run, proviamo a sostituirlo direttamente nel testo del paragrafo
                if placeholder_text in paragraph.text:
                    print(f"Found placeholder in paragraph text: {paragraph.text}")
                    # Salva il testo originale del paragrafo
                    original_text = paragraph.text
                    # Sostituisci solo il segnaposto
                    new_text = original_text.replace(placeholder_text, "trovato")
                    # Aggiorna il testo del paragrafo mantenendo la formattazione
                    replace_in_paragraph_with_formatting(paragraph, original_text, new_text)
                    print("Placeholder replaced with 'trovato' in paragraph text")
                    return
            # Se il valore è "☒" o "☐", sostituisci il segnaposto con il carattere corrispondente
            elif value in ["☒", "☐"]:
                print(f"Using checkbox symbol: {value}")
                for run in paragraph.runs:
                    if placeholder_text in run.text:
                        run.text = run.text.replace(placeholder_text, value)
                return
        # Gestione normale per altri campi
        print("Handling as normal field")
        for run in paragraph.runs:
            if placeholder_text in run.text:
                print(f"Found placeholder in run: {run.text}")
                run.text = run.text.replace(placeholder_text, str(value))
                print(f"Run text after replacement: {run.text}")
    else:
        print(f"Placeholder {placeholder_text} not found in paragraph text")
        # Prova a cercare il segnaposto senza spazi
        placeholder_text_no_spaces = placeholder_text.replace(" ", "")
        if placeholder_text_no_spaces in paragraph.text:
            print(f"Found placeholder without spaces: {placeholder_text_no_spaces}")
            # Gestione speciale per i checkbox
            if placeholder.startswith("checkbox_"):
                print("Found checkbox placeholder")
                if isinstance(value, bool):
                    print(f"Creating checkbox with value: {value}")
                    for run in paragraph.runs:
                        if placeholder_text_no_spaces in run.text:
                            run.text = run.text.replace(placeholder_text_no_spaces, "")
                    create_checkbox_control(paragraph, value)
                    print("Checkbox created successfully")
                    return 